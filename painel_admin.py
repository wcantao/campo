import streamlit as st
import sqlite3
import bcrypt
from collections import defaultdict
from docx import Document
from datetime import datetime
import os

# -----------------------------
# Carregar relatórios por mês e ano
# -----------------------------
def carregar_relatorios_por_mes_ano(mes, ano):
    conn = sqlite3.connect("database.db")
    c = conn.cursor()
    c.execute("""
        SELECT nome, participou, estudos_biblicos, horas, data_envio
        FROM relatorios
        WHERE mes = ? AND ano = ?
        ORDER BY nome, data_envio
    """, (mes, ano))
    data = c.fetchall()
    conn.close()
    return data

# -----------------------------
# Gerar relatório Word
# -----------------------------
def gerar_relatorio_word(mes, ano):
    relatorios = carregar_relatorios_por_mes_ano(mes, ano)
    if not relatorios:
        return None

    doc = Document()
    doc.add_heading(f'Relatório Mensal - {mes}/{ano}', 0)

    agrupado = defaultdict(list)
    for nome, participou, estudos, horas, data_envio in relatorios:
        agrupado[nome].append((participou, estudos, horas, data_envio))

    for nome, registros in agrupado.items():
        doc.add_heading(f'{nome}', level=1)
        for participou, estudos, horas, data_envio in registros:
            p = doc.add_paragraph()
            p.add_run(f'Participou: {"Sim" if participou else "Não"}\n')
            p.add_run(f'Estudos Bíblicos: {estudos}\n')
            p.add_run(f'Horas: {horas}\n')
            p.add_run(f'Data de Envio: {data_envio}\n')
            doc.add_paragraph('-' * 30)

    filename = f"Relatorio_{mes}_{ano}.docx"
    doc.save(filename)
    return filename

# -----------------------------
# Painel admin com filtro por mês e ano, lista bonita + botão Word
# -----------------------------
def mostrar_painel_admin():
    st.title("📊 Painel do Administrador")

    meses = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
             "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    mes_selecionado = st.selectbox("Selecione o mês", meses)
    ano_selecionado = st.number_input("Ano", value=datetime.now().year, step=1, min_value=2000)

    relatorios = carregar_relatorios_por_mes_ano(mes_selecionado, ano_selecionado)
    if not relatorios:
        st.info(f"Nenhum relatório encontrado para {mes_selecionado}/{ano_selecionado}.")
    else:
        agrupado = defaultdict(list)
        for nome, participou, estudos, horas, data_envio in relatorios:
            agrupado[nome].append((participou, estudos, horas, data_envio))

        for nome, registros in agrupado.items():
            st.markdown(f"<h5>👤 {nome}</h5>", unsafe_allow_html=True)
            for idx, (participou, estudos, horas, data_envio) in enumerate(registros, start=1):
                st.markdown(f"""
                    <ul style="font-size: 14px; margin-top: -10px;">
                        <li><b>Participou:</b> {'Sim' if participou else 'Não'}</li>
                        <li><b>Estudos Bíblicos:</b> {estudos}</li>
                        <li><b>Horas:</b> {horas}</li>
                        <li><b>Data de Envio:</b> {data_envio}</li>
                    </ul>
                    <hr style="margin: 5px 0;">
                """, unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("📄 Gerar Relatório Word")
    if st.button("📥 Gerar e Baixar Relatório Word"):
        arquivo = gerar_relatorio_word(mes_selecionado, ano_selecionado)
        if arquivo and os.path.exists(arquivo):
            with open(arquivo, "rb") as f:
                st.download_button("⬇️ Baixar Relatório Word", f, file_name=arquivo)
        else:
            st.warning("Nenhum dado encontrado para este mês/ano.")

# -----------------------------
# Login admin
# -----------------------------
def verificar_admin(usuario, senha):
    conn = sqlite3.connect("database.db")
    c = conn.cursor()
    c.execute('SELECT senha_hash FROM admin WHERE usuario=?', (usuario,))
    result = c.fetchone()
    conn.close()

    if result:
        hash_armazenado = result[0]
        if isinstance(hash_armazenado, str):
            hash_bytes = hash_armazenado.encode()
        else:
            hash_bytes = hash_armazenado

        if bcrypt.checkpw(senha.encode(), hash_bytes):
            return True
    return False

def login_form():
    st.title("🔐 Login do Administrador")
    usuario = st.text_input("Usuário")
    senha = st.text_input("Senha", type="password")

    if st.button("Entrar"):
        if verificar_admin(usuario, senha):
            st.session_state['logado'] = True
        else:
            st.error("Usuário ou senha inválidos!")

# -----------------------------
# Main
# -----------------------------
def main():
    if 'logado' not in st.session_state:
        st.session_state['logado'] = False

    if st.session_state['logado']:
        mostrar_painel_admin()
    else:
        login_form()

if __name__ == "__main__":
    main()





